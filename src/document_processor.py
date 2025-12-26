# src/document_processor.py
from docx import Document
from typing import List, Dict, Tuple
import logging
import re

logger = logging.getLogger(__name__)

class DocumentProcessor:
    def __init__(self, file_path: str):
        self.file_path = file_path
        try:
            self.doc = Document(file_path)
        except Exception as e:
            logger.error(f"Error opening document {file_path}: {e}")
            raise

    def get_paragraphs(self) -> List[str]:
        """Return all paragraphs as a list of text."""
        paragraphs = []
        for p in self.doc.paragraphs:
            text = p.text.strip()
            if text:  # Only add non-empty paragraphs
                paragraphs.append(text)
        return paragraphs

    def get_headings(self) -> List[Dict[str, any]]:
        """Return headings with their level, including implied headings."""
        headings = []
        heading_patterns = [
            (r'^#\s+(.+)$', 1),  # Markdown style # heading
            (r'^##\s+(.+)$', 2), # Markdown style ## heading
            (r'^###\s+(.+)$', 3), # Markdown style ### heading
            (r'^[A-Z][A-Z\s]{5,}:?$', 1),  # ALL CAPS text (likely heading)
            (r'^[IVX]+\.\s+.+$', 1),  # Roman numerals
            (r'^[0-9]+\.\s+.+$', 2),  # Numbered headings
            (r'^[A-Z][^.!?]*:$', 2),  # Text ending with colon
        ]
        
        for p in self.doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
                
            # Check if it's a formal heading style
            style_name = p.style.name
            if style_name.startswith('Heading'):
                try:
                    level = int(style_name.replace('Heading ', ''))
                    headings.append({'text': text, 'level': level})
                    continue
                except ValueError:
                    pass
            
            # Check for implied headings by patterns
            is_heading = False
            level = 2  # Default level for implied headings
            
            # Check if paragraph is short (likely a heading)
            if len(text) < 100 and '\n' not in text:
                # Check for heading patterns
                for pattern, pattern_level in heading_patterns:
                    if re.match(pattern, text, re.IGNORECASE):
                        is_heading = True
                        level = pattern_level
                        break
                
                # Additional heuristic: if text is bold (by checking runs)
                if not is_heading and p.runs:
                    if any(run.bold for run in p.runs):
                        is_heading = True
                        level = 2
                
                # If text is in ALL CAPS and short
                if not is_heading and text.isupper() and len(text.split()) < 10:
                    is_heading = True
                    level = 1
            
            if is_heading:
                headings.append({'text': text, 'level': level})
        
        # If no headings found, use first paragraph as title
        if not headings and self.doc.paragraphs:
            first_para = self.doc.paragraphs[0].text.strip()
            if first_para:
                headings.append({'text': first_para, 'level': 1})
        
        return headings

    def extract_document_structure(self) -> Tuple[List[Dict], List[str]]:
        """
        Extract document structure with headings and their associated paragraphs.
        Returns (headings, sections)
        """
        paragraphs = self.get_paragraphs()

        headings = []
        sections: List[List[str]] = []

        # Heuristics to detect headings are available in get_headings();
        # we'll reuse similar checks inline so we can preserve paragraph order
        heading_patterns = [
            (r'^#\\s+(.+)$', 1),
            (r'^##\\s+(.+)$', 2),
            (r'^###\\s+(.+)$', 3),
            (r'^[A-Z][A-Z\\s]{5,}:?$', 1),
            (r'^[IVX]+\\.\\s+.+$', 1),
            (r'^[0-9]+\\.\\s+.+$', 2),
            (r'^[A-Z][^.!?]*:$', 2),
        ]

        current_section: List[str] = []
        current_heading = None

        for p in self.doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue

            is_heading = False
            level = 2

            # Formal heading style check
            style_name = getattr(p.style, 'name', '')
            if style_name.startswith('Heading'):
                try:
                    level = int(style_name.replace('Heading ', ''))
                except Exception:
                    level = 2
                is_heading = True

            # Pattern based heuristics
            if not is_heading and len(text) < 100 and '\\n' not in text:
                for pattern, pattern_level in heading_patterns:
                    if re.match(pattern, text, re.IGNORECASE):
                        is_heading = True
                        level = pattern_level
                        break

                if not is_heading and p.runs:
                    if any(run.bold for run in p.runs):
                        is_heading = True
                        level = 2

                if not is_heading and text.isupper() and len(text.split()) < 10:
                    is_heading = True
                    level = 1

            if is_heading:
                # Start new heading/section
                if current_heading is not None:
                    sections.append(current_section)
                    current_section = []

                current_heading = {'text': text, 'level': level}
                headings.append(current_heading)
                continue

            # Not a heading: append to current section (or create a default)
            if current_heading is None:
                # No heading seen yet: create an implicit title
                current_heading = {'text': 'Introduction', 'level': 1}
                headings.append(current_heading)
                current_section = []

            current_section.append(text)

        # Append final section
        if current_heading is not None:
            sections.append(current_section)

        # If no headings found at all, create artificial structure
        if not headings:
            return self._create_artificial_structure(paragraphs)

        return headings, sections

    def _create_artificial_structure(self, paragraphs: List[str]) -> Tuple[List[Dict], List[str]]:
        """Create artificial structure when no headings are found."""
        if not paragraphs:
            return [], []
        
        # Use first paragraph as title
        title = paragraphs[0][:50] + "..." if len(paragraphs[0]) > 50 else paragraphs[0]
        headings = [{'text': title, 'level': 1}]
        
        # Group remaining paragraphs into sections
        remaining_paras = paragraphs[1:] if len(paragraphs) > 1 else paragraphs
        
        # Create section headings for groups of paragraphs
        max_paras_per_section = 5
        structured_paras = []
        
        for i in range(0, len(remaining_paras), max_paras_per_section):
            section_paras = remaining_paras[i:i + max_paras_per_section]
            
            # Create a heading for this section
            if i == 0:
                section_heading = "Introduction"
            else:
                section_heading = f"Section {i//max_paras_per_section + 1}"
            
            headings.append({'text': section_heading, 'level': 2})
            structured_paras.append(section_heading)
            structured_paras.extend(section_paras)
        
        return headings, structured_paras

    def debug_document_structure(self):
        """Print debug information about document structure."""
        print("=== DOCUMENT STRUCTURE DEBUG ===")
        print(f"File: {self.file_path}")
        
        paragraphs = self.get_paragraphs()
        print(f"\nTotal paragraphs: {len(paragraphs)}")
        
        for i, para in enumerate(paragraphs):
            print(f"\nParagraph {i+1}:")
            print(f"  Length: {len(para)} chars")
            print(f"  Preview: {para[:100]}...")
        
        headings = self.get_headings()
        print(f"\nTotal headings found: {len(headings)}")
        
        for i, heading in enumerate(headings):
            print(f"\nHeading {i+1}:")
            print(f"  Text: {heading['text']}")
            print(f"  Level: {heading['level']}")
        
        print("\n=== END DEBUG ===")
        return headings, paragraphs